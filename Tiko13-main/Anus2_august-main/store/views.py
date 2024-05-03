from django.shortcuts import render, redirect, get_object_or_404
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView, View, TemplateView
from rest_framework.decorators import api_view, permission_classes, action
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.permissions import IsAuthenticated, IsAuthenticatedOrReadOnly, AllowAny
from datetime import timedelta
from django.utils import timezone
from .models import CommentLike, CommentDislike, ReviewLike, ReviewDislike, Series, Genre, BookUpvote, BookDownvote, Chapter, AuthorNote, Illustration
from django.contrib.auth.models import User, auth
from django.contrib.auth.mixins import LoginRequiredMixin
from .forms import BooksForm, CommentForm, ReviewCreateForm, BookTypeForm, SeriesForm, ChapterForm
from .filters import BookFilter
from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required
from django.http import HttpResponse, HttpResponseRedirect
from django.http import JsonResponse
from django.db.models import Count, Max
from django.urls import reverse_lazy, reverse
from django.views.generic.edit import FormView
from django.contrib import messages
from users.models import Notification, Library, Wallet
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework import generics
from rest_framework.response import Response
from .models import Book, Comment, Review, BookLike, ReviewView, UserBookHistory
from .utils import get_client_ip, is_book_purchased_by_user
from .ip_views import check_book_ip_last_viewed, update_book_ip_last_viewed, check_review_ip_last_viewed, update_review_ip_last_viewed
from .serializer import *
from .converters import create_fb2, parse_fb2
from django.core.files.storage import default_storage
import datetime
from datetime import date
from users.models import WebPageSettings, Library, Profile, FollowersCount, PurchasedBook, NotificationSetting
from django.db.models import Exists, OuterRef
from django.db.models import Q
from rest_framework.exceptions import PermissionDenied
import logging
from django.http import FileResponse, Http404
from .permissions import IsBookAuthor
from users.notification_utils import send_book_update_notifications
from lxml import etree
from docx import Document
import fitz
import zipfile
from bs4 import BeautifulSoup
from django.db import transaction
import filetype
from django.utils.timezone import now


class BooksListAPIView(generics.ListAPIView):
    serializer_class = BookSerializer

    def get_queryset(self):
        user = self.request.user
        base_query = Book.objects.exclude(genre__name='Undefined')  # Exclude books with "Undefined" genre

        if user.is_authenticated:
            return base_query.filter(
                Q(visibility='public') |
                Q(visibility='followers', author__follower_users__follower=user)
            ).distinct().order_by('-views_count')

        return base_query.filter(visibility='public').order_by('-views_count')


class BookDetailAPIView(generics.RetrieveAPIView):
    queryset = Book.objects.all()
    serializer_class = BookSerializer

    def get_object(self):
        book_id = self.kwargs.get('book_id')
        book = get_object_or_404(Book, id=book_id)

        # Allow access to unlisted books only via direct link
        if book.visibility == 'unlisted':
            return book

        # Check if the user has access to the book based on its visibility
        user = self.request.user
        if book.visibility == 'private' and book.author != user:
            raise PermissionDenied('You do not have permission to view this book.')
        elif book.visibility == 'followers' and not book.author.followers.filter(user=user).exists():
            raise PermissionDenied('You do not have permission to view this book.')

        if book.is_adult:
            if not user.is_authenticated:
                raise Http404("This Book is for Adults Only. Please Sign in to Continue.")
            if user.profile.date_of_birth:
                age = (date.today() - user.profile.date_of_birth).days // 365
                if age < 18:
                    raise Http404("This Book is for Adults Only.")
            else:
                raise Http404("Your age is not specified.")

        return book

    def retrieve(self, request, *args, **kwargs):
        instance = self.get_object()

        serialized_data = self.get_serializer(instance).data

        if request.user.is_authenticated:
            # Update user's book reading history if history recording is enabled
            if request.user.profile.record_history:
                UserBookHistory.objects.update_or_create(
                    user=request.user,
                    book=instance,
                    defaults={'last_accessed': timezone.now()}
                )
        else:
            # Handle unauthenticated user scenario, e.g., using sessions
            unlogged_user_history = request.session.get('unlogged_user_history', [])
            if instance.id not in unlogged_user_history:
                unlogged_user_history.append(instance.id)
                if len(unlogged_user_history) > 10:
                    unlogged_user_history.pop(0)
                request.session['unlogged_user_history'] = unlogged_user_history

        # Continue with existing view count and IP tracking logic
        return Response(serialized_data)


@api_view(['GET'])
def get_book_info(request, book_id):
    try:
        book = Book.objects.get(pk=book_id)
    except Book.DoesNotExist:
        return Response({'error': 'Book not found'}, status=404)

    serializer = BookInfoSerializer(book)
    return Response(serializer.data)


@api_view(['GET'])
def get_book_content(request, book_id):
    try:
        book = Book.objects.get(pk=book_id)
    except Book.DoesNotExist:
        return Response({'error': 'Book not found'}, status=404)

    serializer = BookContentSerializer(book)
    return Response(serializer.data)


class BookSearch(ListView):
    template_name = 'store/book_search.html'
    queryset = Book.objects.all()

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['filter'] = BookFilter(self.request.GET, queryset=self.get_queryset())
        return context


class StudioWelcomeAPIView(APIView):
    def post(self, request):
        serializer = BookTypeSerializer(data=request.data)
        if serializer.is_valid():
            book_type = serializer.validated_data.get('book_type')
            new_book = Book.objects.create(author=request.user, book_type=book_type)
            return Response({'message': 'Book created successfully', 'book_id': new_book.id}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class StudioBooksAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        books = Book.objects.filter(author=request.user)
        serializer = StudioBookSerializer(books, context={'request': request}, many=True)
        return Response(serializer.data)

    def patch(self, request, book_id):
        book = Book.objects.get(id=book_id, author=request.user)
        serializer = BookVisibilitySerializer(book, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class StudioSeriesAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Get series created by the user
        series = Series.objects.filter(author=request.user).prefetch_related('books')
        series_data = []
        for serie in series:
            books = serie.books.all()
            books_data = StudioSeriesBooksSerializer(books, many=True, context={'request': request}).data
            series_data.append({
                'id': serie.id,
                'name': serie.name,
                'books': books_data
            })

        # Get books that are not part of any series
        standalone_books = Book.objects.filter(author=request.user, series__isnull=True)
        standalone_books_data = StudioSeriesBooksSerializer(standalone_books, context={'request': request}, many=True).data

        # Combine series and standalone books in the response
        response_data = {
            'series': series_data,
            'standalone_books': standalone_books_data
        }
        return Response(response_data)

    def put(self, request, book_id):
        book = get_object_or_404(Book, id=book_id, author=request.user)
        old_series = book.series
        serializer = StudioBookSerializer(book, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            if 'volume_number' in serializer.validated_data or 'series' in serializer.validated_data:
                self.update_volume_numbers(book, old_series=old_series)
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def update_volume_numbers(self, updated_book, old_series=None):
        # Update volume numbers in the old series if the book was moved from another series
        if old_series:
            old_series_books = old_series.books.exclude(id=updated_book.id).order_by('volume_number')
            for i, book in enumerate(old_series_books, start=1):
                book.volume_number = i
                book.save(update_fields=['volume_number'])

        # Update volume numbers in the new series if the book is part of a series
        if updated_book.series:
            new_series_books = updated_book.series.books.exclude(id=updated_book.id).order_by('volume_number')
            new_books_list = list(new_series_books)
            new_books_list.insert(updated_book.volume_number - 1, updated_book)
            for i, book in enumerate(new_books_list, start=1):
                book.volume_number = i
                book.save(update_fields=['volume_number'])


class StudioCommentsAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Get only top-level comments for books where the logged-in user is the author
        comments = Comment.objects.filter(book__author=request.user, parent_comment__isnull=True).order_by('-timestamp')
        serializer = StudioCommentSerializer(comments, many=True, context={'request': request})
        return Response(serializer.data)

    def post(self, request, book_id):
        book = get_object_or_404(Book, pk=book_id)
        if not book.can_user_comment(request.user):
            return Response({'error': 'You are not allowed to comment on this book.'}, status=status.HTTP_403_FORBIDDEN)

        data = request.data.copy()  # Make a mutable copy of the request data
        data['book'] = book.id  # Ensure the book is correctly associated

        serializer = CreateCommentSerializer(data=data, context={'request': request})
        if serializer.is_valid():
            new_comment = serializer.save(user=request.user)  # User is added here
            return Response(StudioCommentSerializer(new_comment, context={'request': request}).data,
                            status=status.HTTP_201_CREATED)
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, comment_id):
        comment = get_object_or_404(Comment, id=comment_id)
        if comment.book.author != request.user:
            return Response({'error': 'You are not authorized to delete this comment.'},
                            status=status.HTTP_403_FORBIDDEN)

        comment.delete()
        return Response({'status': 'Comment deleted successfully'}, status=status.HTTP_204_NO_CONTENT)


class StudioIllustrationsAPIView(APIView):
    permission_classes = [IsAuthenticated, IsBookAuthor]
    parser_classes = (MultiPartParser, FormParser)

    def get(self, request, *args, **kwargs):
        book_id = kwargs.get('book_id')
        try:
            book = Book.objects.get(id=book_id)
        except Book.DoesNotExist:
            return Response({'error': 'Book not found'}, status=status.HTTP_404_NOT_FOUND)

        # Check if the request user is the author of the book
        self.check_object_permissions(request, book)

        book_serializer = BookSerializer(book, context={'request': request})
        illustrations = Illustration.objects.filter(book=book)
        illustrations_serializer = IllustrationSerializer(illustrations, many=True, context={'request': request})

        return Response({
            'cover_page': book_serializer.data['coverpage'],
            'illustrations': illustrations_serializer.data
        })

    def post(self, request, *args, **kwargs):
        book_id = kwargs.get('book_id')
        try:
            book = Book.objects.get(id=book_id)
        except Book.DoesNotExist:
            return Response({'error': 'Book not found'}, status=status.HTTP_404_NOT_FOUND)

        # Set the book's cover page
        cover_image = request.data.get('cover_image')
        if cover_image:
            book.coverpage = cover_image
            book.save()

        # Add illustrations with descriptions
        for key in request.data.keys():
            if key.startswith('illustration_image_'):
                index = key.split('_')[-1]
                image = request.data[key]
                description = request.data.get(f'illustration_description_{index}', '')

                serializer = IllustrationSerializer(data={'image': image, 'description': description},
                                                    context={'request': request})
                if serializer.is_valid():
                    serializer.save(book=book)
                else:
                    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        return Response({'message': 'Cover page and illustrations updated successfully'},
                        status=status.HTTP_201_CREATED)

    def put(self, request, *args, **kwargs):
        book_id = kwargs.get('book_id')
        illustration_id = kwargs.get('illustration_id')

        try:
            illustration = Illustration.objects.get(id=illustration_id, book__id=book_id)
        except Illustration.DoesNotExist:
            return Response({'error': 'Illustration not found'}, status=status.HTTP_404_NOT_FOUND)

        serializer = IllustrationSerializer(illustration, data=request.data, partial=True, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def patch(self, request, *args, **kwargs):
        book_id = kwargs.get('book_id')
        illustration_id = request.data.get('illustration_id')
        use_for_library_cover = request.data.get('use_for_library_cover', False)

        try:
            book = Book.objects.get(id=book_id)
            illustration = Illustration.objects.get(id=illustration_id, book=book)
        except (Book.DoesNotExist, Illustration.DoesNotExist):
            return Response({'error': 'Book or Illustration not found'}, status=status.HTTP_404_NOT_FOUND)

        # Check if the request user is the author of the book
        self.check_object_permissions(request, book)

        illustration.use_for_library_cover = use_for_library_cover
        illustration.save()

        return Response({'message': 'Illustration updated successfully'}, status=status.HTTP_200_OK)

    def delete(self, request, *args, **kwargs):
        book_id = kwargs.get('book_id')
        illustration_id = kwargs.get('illustration_id')

        try:
            illustration = Illustration.objects.get(id=illustration_id, book__id=book_id)
        except Illustration.DoesNotExist:
            return Response({'error': 'Illustration not found'}, status=status.HTTP_404_NOT_FOUND)

        illustration.delete()
        return Response({'message': 'Illustration deleted successfully'}, status=status.HTTP_204_NO_CONTENT)


class BooksCreateAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        serializer = BookCreateSerializer(data=request.data, context={'request': request})

        if serializer.is_valid():
            book = serializer.save()
            request.session['book_id'] = book.id
            return Response(serializer.data, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class ChapterContentView(APIView):
    def post(self, request, *args, **kwargs):
        serializer = ChapterContentSerializer(data=request.data, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class ChapterListView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id):
        book = Book.objects.get(id=book_id)
        chapters = book.chapters.all().order_by('created')  # Assuming you want them ordered by creation date
        serializer = ChapterSideSerializer(chapters, many=True)
        return Response(serializer.data)

    def post(self, request, book_id):
        serializer = ChapterSerializers(data=request.data, context={'request': request})
        if serializer.is_valid():
            serializer.save(book_id=book_id)  # Ensure the book_id is passed correctly to the save method
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class AddChapterView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id):
        try:
            with transaction.atomic():
                book = Book.objects.select_for_update().get(id=book_id)  # Lock the book record for update
                if book.author != request.user:
                    return Response({'error': 'You are not the author of this book.'}, status=status.HTTP_403_FORBIDDEN)

                # Refresh the book instance to ensure it's up-to-date
                book.refresh_from_db()

                # Check if the book is of type "Short Story / Poem" and already has a chapter
                if book.book_type == 'short_story_poem' and book.chapters.exists():
                    return Response({'error': 'Short Story / Poem books can only have one chapter.'},
                                    status=status.HTTP_400_BAD_REQUEST)

                # Create an empty chapter
                last_chapter_number = book.chapters.aggregate(Max('chapter_number'))['chapter_number__max'] or 0
                next_chapter_number = last_chapter_number + 1

                # Create an empty chapter with the next chapter number
                new_chapter = Chapter.objects.create(book=book, chapter_number=next_chapter_number)
                serializer = ChapterSerializers(new_chapter)

                return Response(serializer.data, status=status.HTTP_201_CREATED)

        except Book.DoesNotExist:
            return Response({'error': 'Book not found.'}, status=status.HTTP_404_NOT_FOUND)


class StudioChapterView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id, chapter_id, format=None):
        try:
            chapter = Chapter.objects.get(book_id=book_id, id=chapter_id)
        except Chapter.DoesNotExist:
            return Response({'error': 'Chapter not found'}, status=status.HTTP_404_NOT_FOUND)

        if chapter.book.author != request.user:
            return Response({'error': 'You are not authorized to view this chapter.'}, status=status.HTTP_403_FORBIDDEN)

        serializer = ChapterContentSerializer(chapter)
        return Response(serializer.data)

    def put(self, request, book_id, chapter_id, format=None):
        try:
            chapter = Chapter.objects.get(book_id=book_id, id=chapter_id)
        except Chapter.DoesNotExist:
            return Response({'error': 'Chapter not found'}, status=status.HTTP_404_NOT_FOUND)

        if chapter.book.author != request.user:
            return Response({'error': 'You do not have permission to edit this chapter'}, status=status.HTTP_403_FORBIDDEN)

        serializer = ChapterContentSerializer(chapter, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, book_id, chapter_id, format=None):
        try:
            chapter = Chapter.objects.get(pk=chapter_id, book_id=book_id)
        except Chapter.DoesNotExist:
            return Response({'error': 'Chapter not found'}, status=status.HTTP_404_NOT_FOUND)

        # Check if the user is authorized to delete the chapter
        if chapter.book.author != request.user:
            return Response({'error': 'Permission denied'}, status=status.HTTP_403_FORBIDDEN)

        chapter.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


@api_view(['POST'])
def toggle_publish_chapter(request, book_id, chapter_id):
    chapter = get_object_or_404(Chapter, id=chapter_id, book__id=book_id)
    if request.user != chapter.book.author:
        return Response({'error': 'You are not authorized to change the publication status of this chapter.'},
                        status=status.HTTP_403_FORBIDDEN)

    # Check if there are any unpublished chapters before this one
    if Chapter.objects.filter(book=book_id, chapter_number__lt=chapter.chapter_number, published=False).exists():
        return Response({'error': 'You cannot publish this chapter until all previous chapters are published.'},
                        status=status.HTTP_403_FORBIDDEN)

    with transaction.atomic():
        chapter.published = not chapter.published
        chapter.save()

        if chapter.published:
            send_book_update_notifications(chapter.book, chapter.title)  # Ensure this function is correctly defined in your utils

    return Response({'published': chapter.published, 'message': 'Chapter publication status toggled successfully.'})


@api_view(['POST'])
def add_author_note(request, book_id, chapter_id):
    try:
        chapter = Chapter.objects.get(book_id=book_id, id=chapter_id)
    except Chapter.DoesNotExist:
        return Response({'error': 'Chapter not found'}, status=404)

    # Data from request
    start = request.data.get('start')
    end = request.data.get('end')
    note_text = request.data.get('note_text')

    # Validate data
    if start is None or end is None or note_text is None:
        return Response({'error': 'Invalid request'}, status=400)

    # Create a new AuthorNote instance
    AuthorNote.objects.create(
        chapter=chapter,
        book=chapter.book,  # Set the book field
        author=request.user,
        start_position=start,
        end_position=end,
        note_text=note_text
    )

    return Response({'message': 'Note added successfully'})


@api_view(['GET'])
def get_author_notes(request, book_id, chapter_id):
    notes = AuthorNote.objects.filter(book_id=book_id, chapter_id=chapter_id, author=request.user)
    serialized_notes = AuthorNoteSerializer(notes, many=True)
    return Response(serialized_notes.data)


class BookNotesView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id):
        notes = AuthorNote.objects.filter(book=book_id, author=request.user)
        serializer = AuthorNoteSerializer(notes, many=True)
        return Response(serializer.data)


class ChapterNotesView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id, chapter_id):
        notes = AuthorNote.objects.filter(book_id=book_id, chapter_id=chapter_id, author=request.user)
        serializer = AuthorNoteSerializer(notes, many=True)
        return Response(serializer.data)


class ChapterDownloadView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id, chapter_id):
        try:
            chapter = Chapter.objects.get(book_id=book_id, id=chapter_id)
        except Chapter.DoesNotExist:
            return Response({'error': 'Chapter not found'}, status=status.HTTP_404_NOT_FOUND)

        if chapter.book.author != request.user:
            return Response({'error': 'Permission denied'}, status=status.HTTP_403_FORBIDDEN)

        fb2_content = create_fb2(chapter)
        response = HttpResponse(fb2_content, content_type='application/xml')
        response['Content-Disposition'] = f'attachment; filename="{chapter.title}.fb2"'
        return response


class ChapterUploadView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id, chapter_id):
        try:
            chapter = Chapter.objects.get(book_id=book_id, id=chapter_id)
        except Chapter.DoesNotExist:
            return Response({'error': 'Chapter not found'}, status=status.HTTP_404_NOT_FOUND)

        if chapter.book.author != request.user:
            return Response({'error': 'Permission denied'}, status=status.HTTP_403_FORBIDDEN)

        file = request.FILES.get('file')
        if not file or not file.name.endswith('.fb2'):
            return Response({'error': 'Invalid file format'}, status=status.HTTP_400_BAD_REQUEST)

        # Save file temporarily
        file_path = default_storage.save(f'tmp/{file.name}', file)
        with default_storage.open(file_path, 'rb') as fb2_file:
            content = parse_fb2(fb2_file)  # Parse FB2 file to extract content

        chapter.content = content
        chapter.save()

        # Cleanup the temporary file
        default_storage.delete(file_path)

        return Response({'message': 'Chapter uploaded successfully'})


class DownloadBookView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id, format):
        book = get_object_or_404(Book, pk=book_id)

        if not book.can_user_download(request.user):
            return Response({'error': 'You are not allowed to download this book.'}, status=status.HTTP_403_FORBIDDEN)

        try:
            # Assuming you have a method to get the file path based on the format
            file_path = book.get_file_path(format)
            return FileResponse(open(file_path, 'rb'), as_attachment=True, filename=f'{book.name}.{format}')
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_404_NOT_FOUND)


class SearchApiView(APIView):
    def get(self, request):
        query = request.query_params.get('q', '')
        books = Book.objects.filter(name__icontains=query)
        book_serializer = BookSerializer(books, many=True)

        # Filter authors who have at least one published book
        authors = User.objects.annotate(
            has_books=Exists(Book.objects.filter(author=OuterRef('pk')))
        ).filter(has_books=True, username__icontains=query).distinct()
        author_serializer = AuthorSerializer(authors, many=True)

        return Response({
            'books': book_serializer.data,
            'authors': author_serializer.data
        })


class BookSettingsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id, format=None):
        book = get_object_or_404(Book, pk=book_id)

        if book.author != request.user:
            return Response({'error': 'You do not have permission to view the settings of this book'}, status=status.HTTP_403_FORBIDDEN)

        # Create the context for the serializer
        followers_usernames = request.user.follower_users.values_list('follower__username', flat=True)
        context = {'co_author_queryset': User.objects.filter(username__in=followers_usernames)}

        serializer = BookSettingsSerializer(book, context=context)
        return Response(serializer.data)

    def patch(self, request, book_id, format=None):
        book = get_object_or_404(Book, pk=book_id)

        if book.author != request.user:
            return Response({'error': 'You do not have permission to edit this book'}, status=status.HTTP_403_FORBIDDEN)

        # Create the context for the serializer
        followers_usernames = request.user.follower_users.values_list('follower__username', flat=True)
        context = {'co_author_queryset': User.objects.filter(username__in=followers_usernames)}

        serializer = BookSettingsSerializer(book, data=request.data, partial=True, context=context)
        if serializer.is_valid():
            if 'is_adult' in serializer.validated_data and serializer.validated_data['is_adult'] is True:
                if not serializer.validated_data.get('confirm_adult_content', False):
                    return Response({'error': 'Please confirm that you want to set this book as adult content.'},
                                    status=status.HTTP_400_BAD_REQUEST)
            serializer.save()
            return Response(serializer.data)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class IllustrationView(APIView):
    def get(self, request, *args, **kwargs):
        illustrations = Illustration.objects.filter(book__id=kwargs['book_id'])
        serializer = IllustrationSerializer(illustrations, many=True, context={'request': request})
        return Response(serializer.data)


class BookSaleView(APIView):
    pass


class BookTextAPIView(APIView):
    """
    Add text to a book.
    """

    def post(self, request, *args, **kwargs):
        book_id = request.session.get('book_id', None)

        if not book_id:
            return Response({"error": "Book ID not found in session."}, status=status.HTTP_400_BAD_REQUEST)

        book = get_object_or_404(Book, id=book_id)
        serializer = ChapterSerializers(data=request.data)

        if serializer.is_valid():
            chapter = serializer.save(book=book)
            return Response(serializer.data, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@method_decorator(staff_member_required, name='dispatch')
class BooksUpdate(UpdateView):
    template_name = 'store/book_create.html'
    form_class = BooksForm
    success_url = '/'
    queryset = Book.objects.all()


@method_decorator(staff_member_required, name='dispatch')
class BooksDelete(DeleteView):
    template_name = 'store/book_delete.html'
    queryset = Book.objects.all()
    success_url = '/'


class SeriesCreateView(LoginRequiredMixin, CreateView):
    model = Series
    form_class = SeriesForm
    template_name = 'store/create_series.html'
    success_url = '/'

    def form_valid(self, form):
        series = form.save(commit=False)
        series.author = self.request.user
        series.save()
        return super().form_valid(form)


class SeriesDetailView(DetailView):
    model = Series
    context_object_name = 'series'
    template_name = 'store/series_detail.html'


class SeriesUpdateView(LoginRequiredMixin, UpdateView):
    model = Series
    form_class = SeriesForm
    template_name = 'store/update_series.html'

    def form_valid(self, form):
        series = form.save(commit=False)
        series.author = self.request.user
        series.save()
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('series_detail', args=[self.object.id])


class CommentListCreateView(APIView):
    def get_permissions(self):
        """
        Instantiates and returns the list of permissions that this view requires.
        """
        if self.request.method == 'POST':
            self.permission_classes = [IsAuthenticated, ]  # Only authenticated users can post comments
        else:
            self.permission_classes = [AllowAny, ]  # Allow anyone to read comments
        return super(CommentListCreateView, self).get_permissions()

    def get(self, request, book_id):
        book = get_object_or_404(Book, pk=book_id)
        comments = Comment.objects.filter(book=book, parent_comment__isnull=True).order_by('-rating')
        serialized_comments = CommentSerializer(comments, many=True, context={'request': request})
        return Response({'comments': serialized_comments.data})

    def post(self, request, book_id):
        book = get_object_or_404(Book, pk=book_id)
        if not book.can_user_comment(request.user):
            return Response({'error': 'You are not allowed to comment on this book.'}, status=status.HTTP_403_FORBIDDEN)

        data = request.data.copy()  # Make a mutable copy of the request data
        data['book'] = book.id  # Ensure the book is correctly associated

        serializer = CreateCommentSerializer(data=data, context={'request': request})
        if serializer.is_valid():
            new_comment = serializer.save(user=request.user)  # User is added here
            return Response(StudioCommentSerializer(new_comment, context={'request': request}).data,
                            status=status.HTTP_201_CREATED)
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class LikeCommentView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id, comment_id):

        comment = get_object_or_404(Comment, id=comment_id, book_id=book_id)
        if comment.user == request.user:
            return Response({'error': 'You cannot like your own comment.'}, status=status.HTTP_403_FORBIDDEN)

        CommentDislike.objects.filter(comment=comment, user=request.user).delete()
        like, created = CommentLike.objects.get_or_create(comment=comment, user=request.user)

        return Response({'status': 'liked' if created else 'like exists'}, status=status.HTTP_200_OK)


class DislikeCommentView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id, comment_id):

        comment = get_object_or_404(Comment, id=comment_id, book_id=book_id)
        if comment.user == request.user:
            return Response({'error': 'You cannot dislike your own comment.'}, status=status.HTTP_403_FORBIDDEN)

        CommentLike.objects.filter(comment=comment, user=request.user).delete()
        dislike, created = CommentDislike.objects.get_or_create(comment=comment, user=request.user)

        return Response({'status': 'disliked' if created else 'dislike exists'}, status=status.HTTP_200_OK)


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def delete_comment(request, comment_id):
    try:
        comment = Comment.objects.get(pk=comment_id, user=request.user)
    except Comment.DoesNotExist:
        return Response({'error': 'Comment not found or not owned by user.'}, status=status.HTTP_404_NOT_FOUND)

    comment.delete()
    return Response({'message': 'Comment deleted successfully.'}, status=status.HTTP_204_NO_CONTENT)


class ReviewListView(generics.ListCreateAPIView):
    serializer_class = ReviewSerializer

    def get_queryset(self):
        book_id = self.kwargs.get('book_id')
        return Review.objects.filter(book_id=book_id)

    def get_serializer_class(self):
        if self.request.method == 'POST':
            return ReviewCreateSerializer
        return ReviewSerializer

    def perform_create(self, serializer):
        serializer.save()


class LikeReviewView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id, review_id):
        review = get_object_or_404(Review, id=review_id, book_id=book_id)
        if review.author == request.user:
            return Response({'error': 'You cannot like your own review.'}, status=status.HTTP_403_FORBIDDEN)

        # Toggle the like
        if request.user in review.likes.all():
            review.likes.remove(request.user)
            status_message = 'like removed'
        else:
            review.likes.add(request.user)
            status_message = 'liked'

        return Response({'status': status_message}, status=status.HTTP_200_OK)



@api_view(['GET'])
def review_detail(request, review_id):
    review = get_object_or_404(Review, id=review_id)
    review.increase_views_count(request)

    serializer = ReviewSerializer(review)
    return Response(serializer.data)


def upvote_book(request, book_id):
    book = get_object_or_404(Book, id=book_id)

    # Check if the user has already downvoted the book
    if BookDownvote.objects.filter(user=request.user, book=book).exists():
        # User has previously downvoted the book, so remove the downvote
        BookDownvote.objects.filter(user=request.user, book=book).delete()

    # Perform the logic for handling the upvote action
    BookUpvote.objects.get_or_create(user=request.user, book=book)

    # Update the book rating
    book.rating = book.upvotes.count() - book.downvotes.count()
    book.save()

    # Redirect back to the same page
    return redirect(reverse('book_detail', kwargs={'pk': book.pk}))


def downvote_book(request, book_id):
    book = get_object_or_404(Book, id=book_id)

    # Check if the user has already upvoted the book
    if BookUpvote.objects.filter(user=request.user, book=book).exists():
        # User has previously upvoted the book, so remove the upvote
        BookUpvote.objects.filter(user=request.user, book=book).delete()

    # Perform the logic for handling the downvote action
    BookDownvote.objects.get_or_create(user=request.user, book=book)

    # Update the book rating
    book.rating = book.upvotes.count() - book.downvotes.count()
    book.save()

    # Redirect back to the same page
    return redirect(reverse('book_detail', kwargs={'pk': book.pk}))


class AddToLikedView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id):
        user = request.user
        book = get_object_or_404(Book, id=book_id)
        library, _ = Library.objects.get_or_create(user=user)

        # Check if the book is adult and the user is under 18
        if book.is_adult:
            if not user.profile.date_of_birth:
                return Response({'error': 'Your age is not specified.'}, status=status.HTTP_400_BAD_REQUEST)
            age = (date.today() - user.profile.date_of_birth).days // 365
            if age < 18:
                return Response({'error': 'This book is for adults only.'}, status=status.HTTP_403_FORBIDDEN)

        # Check if the book is already liked
        if book in library.liked_books.all():
            return Response({'message': 'Book is already liked'}, status=status.HTTP_400_BAD_REQUEST)

        # Add book to the liked category
        library.liked_books.add(book)
        return Response({'message': 'Book added to liked successfully'}, status=status.HTTP_200_OK)



class Reader(APIView):
    permission_classes = [IsAuthenticatedOrReadOnly]

    def get(self, request, book_id):
        try:
            book = Book.objects.get(pk=book_id)
            # Add a check for the book's adult content flag
            if book.is_adult:
                if not request.user.is_authenticated:
                    return Response({'detail': 'Sorry, you need to log in and be older than 18 to read this book.'}, status=status.HTTP_401_UNAUTHORIZED)
                if not self.is_user_adult(request.user):
                    return Response({'detail': 'Sorry, this content is restricted to users over 18.'}, status=status.HTTP_403_FORBIDDEN)

        except Book.DoesNotExist:
            return Response({'detail': 'Book not found.'}, status=status.HTTP_404_NOT_FOUND)

        chapters = Chapter.objects.filter(book=book, published=True)
        user_is_author = request.user == book.author
        user_has_purchased = False
        if request.user.is_authenticated:
            user_has_purchased = book in request.user.library.purchased_books.all()

        can_access_all_chapters = user_is_author or user_has_purchased

        serialized_chapters = self.serialize_chapters(chapters, book, request.user, can_access_all_chapters)

        return Response(serialized_chapters)

    def is_user_adult(self, user):
        try:
            web_page_settings = WebPageSettings.objects.get(profile__user=user)
            dob = web_page_settings.date_of_birth
            if dob is None:
                # If DOB is not set, consider as not adult
                return False
            today = timezone.now().date()
            age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            return age >= 18
        except WebPageSettings.DoesNotExist:
            # Handle cases where the user does not have WebPageSettings or it's not configured properly
            return False

    def serialize_chapters(self, chapters, book, user, user_has_purchased):
        serialized_chapters = []
        for chapter in chapters:
            chapter_data = {
                'id': chapter.id,
                'title': chapter.title,
            }
            if chapter.is_free or user_has_purchased:
                chapter_data['content'] = chapter.content
            else:
                chapter_data['content'] = 'This content is locked. Please purchase the book to read.'
            serialized_chapters.append(chapter_data)

        return serialized_chapters


def is_user_adult(user):
    if not user.is_authenticated:
        return False
    try:
        dob = user.profile.webpagesettings.date_of_birth
        today = timezone.now().date()
        return (today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))) >= 18
    except (AttributeError, WebPageSettings.DoesNotExist):
        return False


class SingleChapterView(APIView):
    def get(self, request, book_id, chapter_id):
        try:
            chapter = Chapter.objects.get(book_id=book_id, id=chapter_id, published=True)
        except Chapter.DoesNotExist:
            return Response({'detail': 'Chapter not found or not published.'}, status=status.HTTP_404_NOT_FOUND)

        is_author = request.user.is_authenticated and chapter.book.author == request.user
        is_adult = is_user_adult(request.user)
        can_access = (
            (chapter.is_free and (not chapter.book.is_adult or is_adult)) or
            (is_book_purchased_by_user(chapter.book, request.user) and (not chapter.book.is_adult or is_adult)) or
            is_author
        )

        if not can_access:
            return Response({'detail': 'You do not have access to this chapter.'}, status=status.HTTP_403_FORBIDDEN)

        serialized_chapter = ChapterSerializers(chapter).data

        # Update user's book reading history if record_history is True
        if request.user.is_authenticated and request.user.profile.record_history:
            UserBookHistory.objects.update_or_create(
                user=request.user,
                book=chapter.book,
                defaults={'last_accessed': timezone.now()}
            )

        if not request.user.is_authenticated:
            history = get_unlogged_user_history(request)
            if len(history) >= 10:
                history.pop(0)  # Remove the oldest entry if history is full
            history.append(chapter.book.id)
            set_unlogged_user_history(request, history)

        return Response(serialized_chapter)


def get_unlogged_user_history(request):
    return request.session.get('unlogged_user_history', [])


def set_unlogged_user_history(request, book_ids):
    request.session['unlogged_user_history'] = book_ids
    request.session.modified = True  # Ensure the session is saved


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def like_book(request, pk):
    book = Book.objects.get(pk=pk)
    BookLike.objects.get_or_create(user=request.user, book=book)
    return Response({'status': 'book liked'})


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def unlike_book(request, pk):
    book = Book.objects.get(pk=pk)
    BookLike.objects.filter(user=request.user, book=book).delete()
    return Response({'status': 'book unliked'})


class PurchaseBookView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id):
        user = request.user
        profile = user.profile
        wallet = get_object_or_404(Wallet, profile=profile)

        book = get_object_or_404(Book, id=book_id)
        book_price = book.price

        if wallet.balance >= book_price:
            wallet.purchase(book, book_price)
            user.library.purchased_books.add(book)
            # Add additional logic if needed, like adding the book to the user's library
            return Response({'message': 'Book purchased successfully'})
        else:
            return Response({'error': 'Insufficient wallet balance'}, status=400)


class RefundBookView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id):
        user = request.user
        book = get_object_or_404(Book, id=book_id)

        try:
            purchased_book = PurchasedBook.objects.get(library=user.library, book=book)
            # Check if the refund is within the allowed time frame (e.g., 30 days)
            if timezone.now() - purchased_book.purchase_date <= timedelta(days=30):
                user.library.purchased_books.remove(book)
                user.profile.wallet.balance += book.price
                user.profile.wallet.save()
                purchased_book.delete()
                return Response({'message': 'Book refunded successfully'})
            else:
                return Response({'error': 'Refund period has expired'}, status=400)
        except PurchasedBook.DoesNotExist:
            return Response({'error': 'Book not found in purchased books'}, status=404)


class HistoryView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        now = timezone.now()
        search_query = request.query_params.get('search', None)
        user_history = UserBookHistory.objects.filter(user=request.user).order_by('-last_accessed')

        if search_query:
            user_history = user_history.filter(book__name__icontains=search_query)

        history_dict = {
            "Today": user_history.filter(last_accessed__date=now.date()),
            "Yesterday": user_history.filter(last_accessed__date=(now - timedelta(days=1)).date()),
            "Last Week": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=7), now.date() - timedelta(days=2)]),
            "Week Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=14), now.date() - timedelta(days=8)]),
            "Two Weeks Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=21), now.date() - timedelta(days=15)]),
            "Three Weeks Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=28), now.date() - timedelta(days=22)]),
            "Month Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=60), now.date() - timedelta(days=29)]),
            "Two Months Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=90), now.date() - timedelta(days=61)]),
            "Three Months Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=180), now.date() - timedelta(days=91)]),
            "Half Year Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=365), now.date() - timedelta(days=181)]),
            "A Year Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=730), now.date() - timedelta(days=366)]),
            "Two Years Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=1095), now.date() - timedelta(days=731)]),
            "A Long Time Ago": user_history.filter(last_accessed__date__lte=now.date() - timedelta(days=1096)),
        }

        # Serialize each queryset and add to response
        for time_category, queryset in history_dict.items():
            history_dict[time_category] = BookViewSerializer(queryset, many=True, context={'request': request}).data

        return Response(history_dict)


def record_history_view(request, book_id):
    if request.user.is_authenticated and request.user.profile.record_history:
        book = get_object_or_404(Book, id=book_id)
        update_user_book_history(request.user, book)
        return Response({'message': 'History recorded successfully'})
    return Response({'error': 'History recording is disabled or user is not authenticated'}, status=status.HTTP_403_FORBIDDEN)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def delete_history(request):
    UserBookHistory.objects.filter(user=request.user).delete()
    return Response({'message': 'History deleted successfully'}, status=status.HTTP_200_OK)


logger = logging.getLogger('user_history')


def update_user_book_history(user, book):
    # Check if there's an existing history entry for this user and book
    history_entry, created = UserBookHistory.objects.get_or_create(user=user, book=book)

    # Update the last_accessed timestamp to the current time
    history_entry.last_accessed = timezone.now()
    history_entry.save()


def update_user_book_history_with_logging(user, book):
    # Update the user book history
    update_user_book_history(user, book)

    # Log the event
    logger.debug(f'Updated history for user {user.username} and book {book.name}')


@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def update_history_settings(request):
    if request.method == 'GET':
        return Response({'record_history': request.user.profile.record_history})

    if request.method == 'POST':
        # Изменяем текущую настройку
        request.user.profile.record_history = not request.user.profile.record_history
        request.user.profile.save()
        return Response({'record_history': request.user.profile.record_history})


class UnloggedUserHistoryView(APIView):
    def get(self, request):
        book_ids = get_unlogged_user_history(request)
        books = Book.objects.filter(id__in=book_ids).order_by('-id')  # Ensure the order reflects the most recent views
        serializer = BookSerializer(books, many=True)
        return Response(serializer.data)

    def post(self, request, book_id):
        book_ids = get_unlogged_user_history(request)

        # Attempt to add the book to history
        if book_id not in book_ids:
            book_ids.append(book_id)
            if len(book_ids) > 10:
                book_ids.pop(0)  # Remove the oldest entry if the limit is exceeded

        set_unlogged_user_history(request, book_ids)
        return Response({'message': 'Book added to history'})


class NewsNotificationsView(APIView):
    def get(self, request):
        # Fetch all notifications for the user, grouped by book
        notifications = Notification.objects.filter(
            recipient=request.user.profile
        ).order_by('-timestamp')

        # Aggregate notifications by book
        books = {}
        for notification in notifications:
            book_id = notification.book_id
            if book_id not in books:
                books[book_id] = {
                    'book': notification.book,
                    'notifications': [],
                    'updates_count': 0
                }
            books[book_id]['notifications'].append(notification)
            books[book_id]['updates_count'] += 1

        # Prepare data for serialization using NewsInfoSerializer
        serialized_data = [
            {
                'book': NewsInfoSerializer(info['book'], context={'request': request}).data,
                'updates_count': info['updates_count'],
                'updates_list': [
                    {
                        'id': n.id,
                        'chapter_title': n.chapter_title,  # Assuming notifications have chapter_title
                        'formatted_timestamp': n.timestamp.strftime('%Y-%m-%d %H:%M:%S')
                    }
                    for n in info['notifications']
                ]
            }
            for info in books.values()
        ]

        return Response(serialized_data)


def parse_fb2_and_create_chapters(file_path, book):
    tree = etree.parse(file_path)
    root = tree.getroot()
    ns = {'fb': 'http://www.gribuser.ru/xml/fictionbook/2.0'}

    last_chapter_number = book.chapters.aggregate(Max('chapter_number'))['chapter_number__max'] or 0

    for section in root.findall('.//fb:section', namespaces=ns):
        title_element = section.find('.//fb:title', namespaces=ns)
        title = "".join(title_element.itertext()) if title_element is not None else None

        content = ""
        for p in section.findall('.//fb:p', namespaces=ns):
            content += "\n".join(p.itertext())

        next_chapter_number = last_chapter_number + 1
        last_chapter_number = next_chapter_number  # Update last chapter number for the next iteration

        Chapter.objects.create(
            book=book,
            title=title,
            content=content.strip(),
            chapter_number=next_chapter_number,  # Добавляем номер главы
            created=timezone.now(),
            updated=timezone.now(),
            published=False
        )


def parse_docx_and_create_chapters(file_path, book):
    doc = Document(file_path)
    for i, _ in enumerate(doc.paragraphs):
        if doc.paragraphs[i].text.startswith('Chapter'):
            title = doc.paragraphs[i].text
            content = '\n'.join(p.text for p in doc.paragraphs[i+1:])
            Chapter.objects.create(book=book, title=title, content=content)
            break


def parse_pdf_and_create_chapters(file_path, book):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    # Если PDF файл представляет собой одну большую главу
    last_chapter_number = book.chapters.aggregate(Max('chapter_number'))['chapter_number__max'] or 0
    next_chapter_number = last_chapter_number + 1

    Chapter.objects.create(
        book=book,
        title="Full Text of PDF",
        content=text,
        chapter_number=next_chapter_number,  # Добавляем номер главы
        created=timezone.now(),
        updated=timezone.now(),
        published=False
    )


def parse_xml_and_create_chapters(file_path, book):
    tree = etree.parse(file_path)
    root = tree.getroot()
    for chapter in root.findall('.//chapter'):
        title = chapter.find('title').text
        content = "".join(p.text for p in chapter.findall('.//paragraph'))
        Chapter.objects.create(book=book, title=title, content=content)


def parse_txt_and_create_chapters(file_path, book):
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
        chapters = text.split('Chapter')  # Предполагаем, что главы разделены словом 'Chapter'

        last_chapter_number = book.chapters.aggregate(Max('chapter_number'))['chapter_number__max'] or 0

        for i, chapter_content in enumerate(chapters):
            if i == 0 and not chapter_content.strip():
                continue  # Пропустить пустую секцию перед первой главой, если таковая есть

            title = f"Chapter {i+1}"
            content = chapter_content.strip()
            next_chapter_number = last_chapter_number + 1
            last_chapter_number = next_chapter_number  # Обновить номер последней главы

            Chapter.objects.create(
                book=book,
                title=title,
                content=content,
                chapter_number=next_chapter_number,
                created=timezone.now(),
                updated=timezone.now(),
                published=False
            )

    if len(chapters) <= 1:
        print("No chapters found in the TXT file.")


def parse_epub_and_create_chapters(file_path, book):
    with zipfile.ZipFile(file_path, 'r') as zf:
        # Предполагаем, что названия файлов глав могут заканчиваться на '.html' или '.xhtml'
        epub_files = [f for f in zf.namelist() if f.endswith(('.html', '.xhtml'))]

        last_chapter_number = book.chapters.aggregate(Max('chapter_number'))['chapter_number__max'] or 0

        for filename in epub_files:
            with zf.open(filename) as file:
                soup = BeautifulSoup(file, 'html.parser')
                title = soup.title.string if soup.title else "Untitled Chapter"

                # Используем метод get_text() для извлечения чистого текста
                content = soup.body.get_text(separator=' ', strip=True) if soup.body else ""

                next_chapter_number = last_chapter_number + 1
                last_chapter_number = next_chapter_number  # Обновляем номер последней главы для следующей итерации

                Chapter.objects.create(
                    book=book,
                    title=title,
                    content=content,
                    chapter_number=next_chapter_number,
                    created=timezone.now(),
                    updated=timezone.now(),
                    published=False
                )

    if not epub_files:
        print("No HTML or XHTML files found in EPUB archive.")


def get_file_type_by_extension(file_path):
    import os
    extension = os.path.splitext(file_path)[1].lower()
    return {
        '.fb2': 'application/x-fictionbook+xml',
        '.epub': 'application/epub+zip',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.txt': 'text/plain',
        '.pdf': 'application/pdf',
        # Добавьте другие расширения и MIME типы
    }.get(extension, None)


class BookFileUploadView(APIView):
    handlers = {
        'application/epub+zip': parse_epub_and_create_chapters,
        'application/pdf': parse_pdf_and_create_chapters,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': parse_docx_and_create_chapters,
        'text/plain': parse_txt_and_create_chapters,
        'application/x-fictionbook+xml': parse_fb2_and_create_chapters,
        # Добавьте другие MIME типы и соответствующие функции обработки
    }

    def post(self, request):
        serializer = BookFileSerializer(data=request.data)
        if serializer.is_valid():
            book = Book.objects.create(author=request.user, book_type='default_type')
            book_file = serializer.save(book=book)
            file_type = get_file_type_by_extension(book_file.file.path)  # Используем новую функцию
            if file_type is None:
                return Response({"error": "Unsupported file type"}, status=status.HTTP_400_BAD_REQUEST)

            handler = self.handlers.get(file_type)
            if handler:
                handler(book_file.file.path, book)
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            else:
                return Response({"error": "No handler for file type"}, status=status.HTTP_400_BAD_REQUEST)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


